# coding=UTF-8
#!/usr/bin/python

from latin import replace_utf8
import codecs

bbl = codecs.open('ieee.bbl','r','utf-8')

class Refer:
    def __init__(self):
        self.author=''
        self.title =''
        self.publish=''
        self.time=''
        self.vol = 0
        self.iss = 0
        self.pp_s = 0  ## pp start
        self.pp_e = 0  ## pp end
        self.year = 0

    def parse_author(self):
        # print('self.author = ',self.author)
        self.author = self.author.replace('\r','')
        self.author = self.author.replace('\n','')
        # self.author = self.author.decode("latex")
        self.author = self.author.replace('{','')
        self.author = self.author.replace('}','')
        self.author = replace_utf8(self.author)

        self.author = self.process_author(self.author)

    def parse_title(self):
        self.title = self.title.replace('\r','')
        self.title = self.title.replace('\n','')
        if self.title.index('{') and self.title.index('}'):
            self.title = self.title[self.title.index('{')+1 : self.title.index('}')]
            # print('title: ' +  self.title)
            # self.title = '“' + self.title + ',” '
            return True
        else:
            return False
    def process_name(self, name):
        name = name.replace('.','')
        out_name = ''
        first_up_get =False
        first_low_get =  False # for this Leemon~C Baird~III
        for c in name[::-1]:
            if c.islower():
                first_low_get = True
            if c.isupper(): 
                if not first_up_get:
                    out_name = ' ' + c + out_name
                else:
                    out_name = c + '.' + out_name

                first_up_get = True if first_low_get else False
            else:
                if not first_up_get:
                    out_name = c + out_name

        return out_name

    def process_author(self, authors):
        autor_ary = authors.split('‚')
        # print('autor_ary',autor_ary)
        new_authors = ''
        if len(autor_ary) > 1:
            for i, name in enumerate(autor_ary):
                out_name = self.process_name(name)
                new_authors += out_name 

                new_authors += ', ' if i < len(autor_ary)-2 else ', and ' if i < len(autor_ary)-1 else ''
        else:
            # 2 authors
            autor_ary = authors.split(' and')
            if len(autor_ary) > 1:
                new_authors = self.process_name(autor_ary[0]) +' and ' + self.process_name(autor_ary[1])
            else:
                new_authors = self.process_name(autor_ary[0])
            
            
        # print('new_authors',new_authors)
        return new_authors

    def parse_publish(self):
        self.publish = self.publish.replace('\r','')
        self.publish = self.publish.replace('\n','')
        self.publish = self.publish.replace('In','in')
        self.publish = self.publish.replace('\\newblock','')
        # print('publish: ' +  self.publish)
        
        if '\\penalty0' in self.publish:
            self.parse_vol_iss_pp(self.publish)

        if '\emph{' in self.publish  and '}' in self.publish :
            # self.publish = self.publish[self.publish.index('{')+1 : self.publish.index('}')]
            # self.publish = self.publish.replace('\\emph{','<i>')
            # self.publish = self.publish.replace('}','</i>')
            self.year = self.publish[self.publish.rfind(',')+1: self.publish.rfind('.') ]
            
            self.publish = self.publish[    \
                        self.publish.index('\\emph{')+6: self.publish.index('}') ]
            
            
        # print('publish: ' +  self.publish)

    # \newblock \emph{IEEE Trans. on Pattern Analysis and Machine Intelligence},
    #       35\penalty0 (8):\penalty0 1798--1828, 2013.
    #  vol. 36, iss.8, pp. 1798-1828
    def parse_vol_iss_pp(self, p):
        # print('-----parse_vol-----')
        comma_s = p.rfind(',', 0, p.index('\\penalty0') )
        comma_e = p.rfind(',', comma_s)
        vol_iss_pp = p[comma_s+1:comma_e]

        # print("vol_iss_pp={}".format(vol_iss_pp))
        self.vol = vol_iss_pp[:vol_iss_pp.index('\\penalty0')]
        self.vol = self.vol.replace(' ','')
        vol_iss_pp = vol_iss_pp[vol_iss_pp.index('\\penalty0')+9:]

        # print('after vol -> '+ str(vol_iss_pp))
        if '(' in vol_iss_pp and ')' in vol_iss_pp:
            self.iss = vol_iss_pp[vol_iss_pp.index('(') + 1 : vol_iss_pp.index(')') ] 
            vol_iss_pp =  vol_iss_pp[vol_iss_pp.index(')')+1:]

        # print('after iss -> '+ str(vol_iss_pp))

        # vol_iss_pp = vol_iss_pp.replace('\\:penalty0','')
        if '--' in vol_iss_pp:
            self.pp_s = vol_iss_pp[vol_iss_pp.index(' ')+1: vol_iss_pp.index('--')]
            self.pp_e = vol_iss_pp[vol_iss_pp.index('--')+2: ]

        # print('vol={}, iss={}, p_s={}, p_e={} '.format(self.vol, self.iss, self.pp_s, self.pp_e))


M_NONE    = 0
M_BIBITEM = 1
M_AUTHOR  = 2
M_TITLE   = 3
M_PUBLISH = 4




bibitem_block = False
read_author = False

mode = M_NONE
refer_ary = []

lines = bbl.readlines()

print(len(lines))

for idx, line in enumerate(lines):
    line = line.replace('{\\natexlab{a}}','')
    line = line.replace('{\\natexlab{b}}','')
            
    # print("(%2d)[%3d]:%s " %(idx+1, len(line), line))
    if line.startswith('\\bibitem'):
        # print('in bibitem')
        mode = M_BIBITEM
        now_refer = Refer()
    
    if mode == M_BIBITEM:
        if '}' in line and not '{\\' in line:
            mode = M_AUTHOR
    elif mode == M_AUTHOR:
        # print('in author')
        if not line.startswith('\\newblock'):
            now_refer.author += line
        else:
            now_refer.parse_author()
            # print('author: ' + now_refer.author)
            now_refer.title = line
            
            mode = M_TITLE
    elif mode == M_TITLE:
        if not line.startswith('\\newblock'):
            now_refer.title += line
        else:
            # print('title: ' + now_refer.title)
            now_refer.parse_title()
            now_refer.publish = line
            mode = M_PUBLISH
    elif mode == M_PUBLISH:
        if len(line)>1:
            now_refer.publish += line
        else:
            now_refer.parse_publish()
            refer_ary.append(now_refer)
            mode = M_NONE


from docx import Document
from docx.shared import Inches

document = Document()


p = document.add_paragraph('')
for i, r in enumerate(refer_ary):
    # s = '[{}] {} ,“{},” <i>{}</i>'.format(i+1, r.author, r.title, r.publish)
    s = '[{}] {} ,“{},” '.format(i+1, r.author, r.title, r.publish)
    # p = document.add_paragraph(s)
    p.add_run(s)
    p.add_run(r.publish).italic = True        

    s = ''
    s = (s + ', vol.' + r.vol) if r.vol and len(r.vol) > 0 else s
    s = (s + ', iss.' + r.iss) if r.iss else s
    s = (s + ', pp.{}-{}'.format(r.pp_s, r.pp_e)) if r.pp_s else s
    s = (s + ',' + str(r.year)) if r.year else s

    p.add_run(s) 
    p.add_run('\n')



document.save('refer.docx')
print('Output to refer.docx')